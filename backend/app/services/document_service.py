from typing import Dict, Any, Optional
import io
import logging
import re
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
import httpx
from bs4 import BeautifulSoup
import base64

logger = logging.getLogger(__name__)


class DocumentService:
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md", ".xlsx", ".xls", ".csv", ".png", ".jpg", ".jpeg", ".webp"}
    MAX_FILE_SIZE_MB = 10
    
    def extract_text_from_pdf(self, file_content: bytes) -> str:
        
        try:
            
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for page_num, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {page_num + 1}]\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract page {page_num + 1}: {e}")
                    continue
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise ValueError(f"Could not extract text from PDF: {e}")
    
    def extract_text_from_docx(self, file_content: bytes) -> str:
       
        try:
            
            
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise ValueError(f"Could not extract text from DOCX: {e}")
    
    def extract_text_from_txt(self, file_content: bytes) -> str:
       
        try:
            # Try different encodings
            for encoding in ["utf-8", "latin-1", "cp1252"]:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # Fallback with error replacement
            return file_content.decode("utf-8", errors="replace")
        except Exception as e:
            logger.error(f"TXT extraction failed: {e}")
            raise ValueError(f"Could not read text file: {e}")
    
    def extract_text_from_excel(self, file_content: bytes) -> str:
        
        try:
            excel_file = io.BytesIO(file_content)
            # Read all sheets
            dfs = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
            
            text_parts = []
            for sheet_name, df in dfs.items():
                text_parts.append(f"--- Sheet: {sheet_name} ---")
                # Convert to markdown or string for better readability
                text_parts.append(df.to_string(index=False))
            
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Excel extraction failed: {e}")
            raise ValueError(f"Could not extract text from Excel: {e}")

    def extract_text_from_csv(self, file_content: bytes) -> str:
        
        try:
            csv_file = io.BytesIO(file_content)
            df = pd.read_csv(csv_file)
            return df.to_string(index=False)
        except Exception as e:
            logger.error(f"CSV extraction failed: {e}")
            raise ValueError(f"Could not extract text from CSV: {e}")

    def _get_api_key(self, db, env_name: str, db_key: str) -> str:
        try:
            from app.services.config_service import config_service
            key = config_service.get(db_key)
            if key and isinstance(key, str) and key.strip():
                return key
        except Exception as e:
            logger.warning(f"Error fetching API key for {env_name}: {e}")
        return ""

    async def _analyze_image_with_vlm(self, file_content: bytes, filename: str, db=None) -> str:
        logger.info(f"Analyzing image {filename} using centralized AIExecutionService...")
        
        # Determine mime type
        mime_type = "image/png"
        if filename.lower().endswith((".jpg", ".jpeg")):
            mime_type = "image/jpeg"
        elif filename.lower().endswith(".webp"):
            mime_type = "image/webp"

        prompt = (
            "Describe this image in detail. Extract any text, labels, numbers, tables, or charts visible in the image. "
            "Explain the main contents and visual elements clearly so it can be indexed for search."
        )

        from app.services.ai.execution_service import AIExecutionService, AIFeatureRegistry, current_execution_context
        ctx = current_execution_context.get()
        user_id = ctx.user_id if ctx else "system"

        result = await AIExecutionService.execute(
            db=db,
            workspace_id="system_workspace",
            user_id=user_id,
            feature_key=AIFeatureRegistry.CHAT,
            prompt=prompt,
            model="auto",
            media_data=file_content,
            mime_type=mime_type,
            bypass_billing=True
        )

        text = result.get("text", "")
        if not text:
            raise ValueError("VLM image analysis returned empty response")

        return text

    async def process_file(self, file_content: bytes, filename: str, db: Optional[Any] = None) -> Dict[str, Any]:
       
        # Validate file size
        size_mb = len(file_content) / (1024 * 1024)
        if size_mb > self.MAX_FILE_SIZE_MB:
            raise ValueError(f"File too large: {size_mb:.1f}MB. Maximum: {self.MAX_FILE_SIZE_MB}MB")
        
        # Determine file type
        filename_lower = filename.lower()
        
        if filename_lower.endswith(".pdf"):
            text = self.extract_text_from_pdf(file_content)
            content_type = "pdf"
        elif filename_lower.endswith((".docx", ".doc")):
            text = self.extract_text_from_docx(file_content)
            content_type = "docx"
        elif filename_lower.endswith((".txt", ".md")):
            text = self.extract_text_from_txt(file_content)
            content_type = "txt"
        elif filename_lower.endswith((".xlsx", ".xls")):
            text = self.extract_text_from_excel(file_content)
            content_type = "excel"
        elif filename_lower.endswith(".csv"):
            text = self.extract_text_from_csv(file_content)
            content_type = "csv"
        elif filename_lower.endswith((".png", ".jpg", ".jpeg", ".webp")):
             text = await self._analyze_image_with_vlm(file_content, filename, db)
             content_type = "image"
        else:
            raise ValueError(f"Unsupported file type: {filename}")
        
        if not text or len(text.strip()) < 10:
            raise ValueError("No text content could be extracted from the file")
        
        return {
            "text": text,
            "filename": filename,
            "content_type": content_type,
            "size_bytes": len(file_content),
            "char_count": len(text),
            "word_count": len(text.split())
        }


class URLScraperService:
   
    TIMEOUT_SECONDS = 30
    MAX_CONTENT_LENGTH = 5 * 1024 * 1024  # 5MB
    
    async def scrape_url(self, url: str) -> Dict[str, Any]:
       
           
        # Validate URL
        if not url.startswith(("http://", "https://")):
            raise ValueError("URL must start with http:// or https://")
        
        try:
            async with httpx.AsyncClient(
                timeout=self.TIMEOUT_SECONDS,
                follow_redirects=True,
                headers={"User-Agent": "AuromindAI/1.0 (Knowledge Indexer)"}
            ) as client:
                response = await client.get(url)
                response.raise_for_status()
                
                # Check content length
                if len(response.content) > self.MAX_CONTENT_LENGTH:
                    raise ValueError("Page content too large")
                
                # Parse HTML
                soup = BeautifulSoup(response.content, "lxml")
                
                # Extract title
                title = soup.title.string if soup.title else url
                
                # Remove script and style elements
                for element in soup(["script", "style", "nav", "footer", "header", "aside"]):
                    element.decompose()
                
                # Extract main content
                main_content = soup.find("main") or soup.find("article") or soup.find("body")
                
                if main_content:
                    text = self._extract_text(main_content)
                else:
                    text = soup.get_text(separator="\n", strip=True)
                
                # Clean text
                text = self._clean_text(text)
                
                if not text or len(text.strip()) < 50:
                    raise ValueError("No meaningful content found on page")
                
                return {
                    "text": text,
                    "title": title.strip() if title else url,
                    "url": url,
                    "content_type": "url",
                    "char_count": len(text),
                    "word_count": len(text.split())
                }
                
        except httpx.RequestError as e:
            logger.error(f"Request failed for {url}: {e}")
            raise ValueError(f"Failed to fetch URL: {e}")
        except Exception as e:
            logger.error(f"Scraping failed for {url}: {e}")
            raise ValueError(f"Failed to scrape URL: {e}")
    
    def _extract_text(self, element) -> str:
      
        texts = []
        for child in element.descendants:
            if child.name in ["p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "td", "th"]:
                text = child.get_text(strip=True)
                if text:
                    texts.append(text)
        return "\n\n".join(texts)
    
    def _clean_text(self, text: str) -> str:
       
        # Remove excessive whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' {2,}', ' ', text)
        # Remove common boilerplate patterns
        text = re.sub(r'Cookie[s]? (policy|consent|notice).*?(\n|$)', '', text, flags=re.IGNORECASE)
        return text.strip()


# Global instances
_document_service: Optional[DocumentService] = None
_url_scraper: Optional[URLScraperService] = None


def get_document_service() -> DocumentService:
    
    global _document_service
    if _document_service is None:
        _document_service = DocumentService()
    return _document_service


def get_url_scraper() -> URLScraperService:
    
    global _url_scraper
    if _url_scraper is None:
        _url_scraper = URLScraperService()
    return _url_scraper
